"""Concierge Master â€” Streamlit + Supabase.

InstalaciÃ³n:
    pip install -r requirements_streamlit.txt
EjecuciÃ³n:
    streamlit run concierge_master_app.py

Secrets requeridos en .streamlit/secrets.toml:
    SUPABASE_URL = "..."
    SUPABASE_KEY = "..."
    DELETE_PASSWORD = "..."  # opcional, pero recomendado
"""

from __future__ import annotations

import html
import os
from datetime import datetime, timedelta
from io import BytesIO
from urllib.parse import urlencode

import pandas as pd
import streamlit as st
from supabase import Client, create_client

try:
    from st_aggrid import AgGrid, GridOptionsBuilder, GridUpdateMode, JsCode
except ImportError:
    st.error(
        "Falta la dependencia `streamlit-aggrid`. Ejecuta: "
        "`pip install streamlit-aggrid`"
    )
    st.stop()


# -----------------------------------------------------------------------------
# ConfiguraciÃ³n
# -----------------------------------------------------------------------------

st.set_page_config(
    page_title="Concierge Master",
    page_icon="âœ¦",
    layout="wide",
    initial_sidebar_state="collapsed",
)

TABLE_NAME = "huespedes"
DB_COLUMNS = [
    "id", "eta", "name", "qty", "room", "email", "check_in", "check_out",
    "res_number", "phone", "info", "ird", "hsk", "rate", "trans",
]
DISPLAY_COLUMNS = [
    "eta", "name", "qty", "room", "email", "check_in", "check_out", "nights",
    "res_number", "phone", "info", "ird", "hsk", "rate", "trans",
]
IMPORT_COLUMNS = [column for column in DB_COLUMNS if column != "id"]

CATEGORY_COLORS = {
    "VIP": "#00E5FF",
    "BIRTHDAY": "#FF5252",
    "HONEYMOON": "#FF9800",
    "BABYMOON": "#A78BFA",
    "ANNIVERSARY": "#4ADE80",
    "RELAXURY": "#F472B6",
    "TEAM MEMBER": "#FACC15",
}

# Los enlaces son URL proporcionadas por el usuario. Abren una pestaÃ±a nueva.
QUICK_LINKS = [
    (
        "ALICE",
        "https://auth.aliceapp.com/login-staff?__hstc=85647430.18528c557a8d4857356bbdc77be22153.1745273864718.1745273864718.1745273864718.1&__hssc=85647430.2.1745273864718&__hsfp=92250610",
        "#6C5CE7",
    ),
    (
        "ARRIVALS",
        "https://hilton-my.sharepoint.com/shared?listurl=https%3A%2F%2Fhilton%2Dmy%2Esharepoint%2Ecom%2Fpersonal%2Fefrem%5Fcatellani%5Fwaldorfastoria%5Fcom%2FDocuments&e=5%3A5760d5a1b59d4b69adb09d888a758bb4&sharingv2=true&fromShare=true&at=9&CT=1782844742090&OR=OWA%2DNT%2DMail&SI=NonSentItems&clickParams=eyJYLUFwcE5hbWUiOiJNaWNyb3NvZnQgT3V0bG9vayBXZWIgQXBwIiwiWC1BcHBWZXJzaW9uIjoiMjAyNjA2MTkwMTAuMTIiLCJPUyI6IldpbmRvd3MgMTEifQ%3D%3D&cidOR=Client&id=%2Fpersonal%2Fefrem%5Fcatellani%5Fwaldorfastoria%5Fcom%2FDocuments%2FARRIVAL%20DAYS%2F2026&FolderCTID=0x0120000A5710A5FF38F342BA540726A6B97804",
        "#0284C7",
    ),
    ("LA CENIA", "https://lacerniaadventures.com/", "#059669"),
    ("NO LIMIT", "https://www.experiencecollectioncr.com/", "#EA580C"),
    ("OPEN TABLE", "https://guestcenter.opentable.com/login", "#DC2626"),
    (
        "OUTLOOK-FW",
        "https://outlook.office365.com/mail/inbox/id/AAQkAGMyMWEwZDZkLTk2NDQtNDZiMC1hMmE1LWIxYjFmZGJjYjBmOAAQAIbRdEConWFGtPTirYcPWFY%3D",
        "#2563EB",
    ),
    (
        "OUTLOOK-PC",
        "https://outlook.cloud.microsoft/mail/personalconcierge.costarica@waldorfastoria.com/",
        "#0078D4",
    ),
    ("RELAXURY", "https://relaxury.agilesd.com/", "#DB2777"),
    (
        "VTC",
        "https://hilton-my.sharepoint.com/shared?listurl=https%3A%2F%2Fhilton%2Dmy%2Esharepoint%2Ecom%2Fpersonal%2Fefrem%5Fcatellani%5Fwaldorfastoria%5Fcom%2FDocuments&e=5%3A8e5918d1d45b4d6b90289e3f445b4d82&sharingv2=true&fromShare=true&at=9&cidOR=SPO&id=%2Fpersonal%2Fefrem%5Fcatellani%5Fwaldorfastoria%5Fcom%2FDocuments%2FVIRTUAL%20SIGNATURES&FolderCTID=0x0120000A5710A5FF38F342BA540726A6B97804",
        "#D97706",
    ),
]


# -----------------------------------------------------------------------------
# Estilos
# -----------------------------------------------------------------------------

st.markdown(
    """
<style>
    :root { color-scheme: dark; }
    .stApp { background: #080b12; color: #edf8ff; }
    header[data-testid="stHeader"] { display: none; }
    .block-container { max-width: 100%; padding: .7rem 1.1rem 1.4rem; }
    [data-testid="stVerticalBlock"] { gap: .45rem; }
    [data-testid="stTextInput"] input, [data-testid="stDateInput"] input {
        color: #effaff !important; background: #101827 !important;
        border: 1px solid #263b53 !important; border-radius: 8px !important;
    }
    [data-testid="stTextInput"] input:focus, [data-testid="stDateInput"] input:focus {
        border-color: #00e5ff !important; box-shadow: 0 0 0 1px #00e5ff !important;
    }
    div[data-testid="stMetric"] {
        border: 1px solid #1e3348; background: #0d1420; border-radius: 10px; padding: 10px;
    }
    div[data-testid="stMetric"] label { color: #8ca4ba !important; }
    div[data-testid="stMetric"] [data-testid="stMetricValue"] { color: #00e5ff !important; }
    .quick-links { display: grid; grid-template-columns: repeat(9, minmax(88px, 1fr)); gap: 7px; }
    .quick-link, .action-link {
        display: flex; align-items: center; justify-content: center; min-height: 32px;
        padding: 6px 9px; border-radius: 7px; color: white !important; text-decoration: none !important;
        font: 800 11px/1.1 "Segoe UI", sans-serif; letter-spacing: .25px; text-align: center;
        box-shadow: inset 0 0 0 1px rgba(255,255,255,.15), 0 3px 9px rgba(0,0,0,.25);
        transition: transform .12s ease, filter .12s ease;
    }
    .quick-link:hover, .action-link:hover { filter: brightness(1.15); transform: translateY(-1px); }
    .action-links { display: grid; grid-template-columns: repeat(6, minmax(100px, 1fr)); gap: 8px; margin: 10px 0 3px; }
    .action-link { min-height: 36px; font-size: 11px; }
    .page-title { display:flex; justify-content:space-between; gap:18px; align-items:center; margin-bottom:10px; }
    .brand { display:flex; align-items:center; gap:11px; }
    .brand-mark { color:#d4af37; font: 32px Georgia, serif; letter-spacing:5px; padding-right:11px; border-right:1px solid #766129; }
    .brand-name { color:#d4af37; font-size:11px; font-weight:800; letter-spacing:1.7px; }
    .brand-place { color:#6f879a; font-size:9px; letter-spacing:1px; margin-top:2px; }
    .brand-product { color:#00e5ff; font-size:14px; font-weight:800; margin-top:3px; }
    .clock { color:#00e5ff; font-size:16px; font-weight:800; text-align:right; text-shadow:0 0 12px rgba(0,229,255,.45); }
    .clock-date { color:#00e5ff; font-size:14px; font-weight:700; margin-top:4px; }
    .panel { background:#0d1420; border:1px solid #1e3348; border-radius:10px; padding:12px; }
    .panel-title { color:#00e5ff; font-size:11px; font-weight:800; letter-spacing:.7px; text-transform:uppercase; margin-bottom:8px; }
    .total-strip { border:1px solid #00e5ff; color:#dffcff; background:linear-gradient(90deg,#0b2130,#0d1420); border-radius:8px; padding:7px 12px; text-align:center; font-size:12px; }
    .total-strip strong { color:#00e5ff; font-size:17px; margin-left:6px; }
    .category-row { display:flex; align-items:center; gap:8px; margin:6px 0; }
    .category-label { color:#b9cad8; font-size:10px; width:92px; text-align:right; white-space:nowrap; }
    .category-track { flex:1; height:13px; background:#172331; border-radius:5px; overflow:hidden; }
    .category-fill { height:100%; border-radius:5px; }
    .category-value { color:#e8f8ff; font-size:11px; font-weight:700; width:22px; }
    .selection-banner { margin: 8px 0; padding:8px 11px; background:#062233; border:1px solid #00e5ff; color:#edfaff; border-radius:8px; font-size:12px; }
    .selection-banner b { color:#00e5ff; }
    .stButton > button, .stDownloadButton > button, [data-testid="stFormSubmitButton"] > button {
        background:#1c2b3a !important; color:#eefaff !important; border:1px solid #35526d !important;
        border-radius:7px !important; font-weight:750 !important;
    }
    .stButton > button:hover, .stDownloadButton > button:hover, [data-testid="stFormSubmitButton"] > button:hover {
        border-color:#00e5ff !important; color:#00e5ff !important;
    }
    @media (max-width: 980px) {
        .quick-links { grid-template-columns: repeat(3, 1fr); }
        .action-links { grid-template-columns: repeat(2, 1fr); }
        .clock { display:none; }
    }
</style>
    """,
    unsafe_allow_html=True,
)


# -----------------------------------------------------------------------------
# Helpers
# -----------------------------------------------------------------------------

def safe_text(value: object) -> str:
    """Escapa datos de base de datos antes de interpolarlos en HTML."""
    if value is None or pd.isna(value):
        return ""
    return html.escape(str(value))


def parse_fecha(value: object) -> datetime | None:
    if value is None or pd.isna(value) or not str(value).strip():
        return None
    value = str(value).strip()
    for pattern in ("%B %d, %Y", "%b %d, %Y", "%Y-%m-%d", "%B %d", "%b %d"):
        try:
            return datetime.strptime(value, pattern)
        except ValueError:
            continue
    return None


def normalizar_fecha(value: object) -> str:
    date = parse_fecha(value)
    if date:
        if date.year == 1900:
            date = date.replace(year=datetime.now().year)
        return date.strftime("%B %d, %Y")
    return "" if value is None or pd.isna(value) else str(value).strip()


def calcular_noches(check_in: object, check_out: object) -> int | None:
    """Devuelve checkout - checkin, sin contar una noche negativa."""
    start, end = parse_fecha(check_in), parse_fecha(check_out)
    if not start or not end:
        return None
    current_year = datetime.now().year
    if start.year == 1900:
        start = start.replace(year=current_year)
    if end.year == 1900:
        end = end.replace(year=current_year)
    nights = (end - start).days
    return nights if nights >= 0 else None


def date_from_filter(value: str | None) -> datetime | None:
    try:
        return datetime.strptime(value or "", "%Y-%m-%d")
    except ValueError:
        return None


def url_with(**params: str) -> str:
    clean = {key: value for key, value in params.items() if value not in (None, "")}
    return "?" + urlencode(clean) if clean else "?"


def clear_selection() -> None:
    for key in ("selected_reservation", "selected_reservation_id"):
        st.session_state.pop(key, None)


def get_action() -> str:
    return str(st.query_params.get("action", ""))


def set_action(action: str) -> None:
    st.query_params["action"] = action
    st.rerun()


def clear_page() -> None:
    clear_selection()
    st.query_params.clear()
    st.rerun()


# -----------------------------------------------------------------------------
# Supabase CRUD
# -----------------------------------------------------------------------------

@st.cache_resource
def init_supabase() -> Client:
    return create_client(st.secrets["SUPABASE_URL"], st.secrets["SUPABASE_KEY"])


try:
    supabase = init_supabase()
except KeyError:
    st.error("Configura `SUPABASE_URL` y `SUPABASE_KEY` en los Secrets de Streamlit.")
    st.stop()


@st.cache_data(ttl=30, show_spinner=False)
def cargar_reservaciones() -> pd.DataFrame:
    response = supabase.table(TABLE_NAME).select("*").execute()
    df = pd.DataFrame(response.data)
    if df.empty:
        return pd.DataFrame(columns=DB_COLUMNS + ["nights"])

    for column in DB_COLUMNS:
        if column not in df.columns:
            df[column] = "" if column != "qty" else 0

    df["nights"] = df.apply(
        lambda row: calcular_noches(row.get("check_in"), row.get("check_out")), axis=1
    )
    sort_date = df["check_in"].map(parse_fecha)
    return df.assign(_sort_date=sort_date).sort_values(
        by=["_sort_date", "name"], na_position="last"
    ).drop(columns="_sort_date")


def insertar_reserva(data: dict) -> None:
    supabase.table(TABLE_NAME).insert(data).execute()
    st.cache_data.clear()


def actualizar_reserva(reservation_id: object, data: dict) -> None:
    supabase.table(TABLE_NAME).update(data).eq("id", reservation_id).execute()
    st.cache_data.clear()


def eliminar_reserva(reservation_id: object) -> None:
    supabase.table(TABLE_NAME).delete().eq("id", reservation_id).execute()
    st.cache_data.clear()


def insertar_lote(records: list[dict]) -> None:
    if records:
        supabase.table(TABLE_NAME).insert(records).execute()
        st.cache_data.clear()


# -----------------------------------------------------------------------------
# Exportaciones
# -----------------------------------------------------------------------------

def exportar_excel_por_categorias(df: pd.DataFrame) -> BytesIO:
    from openpyxl import Workbook
    from openpyxl.styles import Alignment, Border, Font, PatternFill, Side

    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Arrivals"

    fill_section = PatternFill("solid", fgColor="00B0F0")
    fill_header = PatternFill("solid", fgColor="123047")
    fill_data = PatternFill("solid", fgColor="F4F7F9")
    white_bold = Font(name="Calibri", size=11, bold=True, color="FFFFFF")
    black_font = Font(name="Calibri", size=10, color="000000")
    center = Alignment(horizontal="center", vertical="center", wrap_text=True)
    left = Alignment(horizontal="left", vertical="center", wrap_text=True)
    border = Border(*(Side(style="thin", color="D6DEE4") for _ in range(4)))

    groups = [
        ("CUMPLEAÃ‘OS", ("BIRTHDAY", "CUMPLE", "BDAY")),
        ("VIP", ("VIP",)),
        ("HONEYMOON", ("HONEYMOON", "LUNA DE MIEL")),
        ("ANNIVERSARY", ("ANNIVERSARY", "ANIVERSARIO")),
        ("BABYMOON", ("BABYMOON",)),
        ("TEAM MEMBER", ("TEAM MEMBER", "STAFF", "EMPLOYEE")),
        ("GENERAL", ()),
    ]
    export_columns = [
        ("id", "ID"), ("eta", "ETA"), ("name", "NAME"), ("qty", "QTY"),
        ("room", "ROOM"), ("email", "EMAIL"), ("check_in", "CHECK IN"),
        ("check_out", "CHECK OUT"), ("nights", "NIGHTS"),
        ("res_number", "RESERVATION"), ("phone", "PHONE"), ("info", "INFORMATION"),
        ("ird", "IRD"), ("hsk", "HSK"), ("rate", "RATE"), ("trans", "TRANSPORTATION"),
    ]

    remaining = df.copy()
    row_number = 1
    for title, keywords in groups:
        if keywords:
            mask = remaining["info"].fillna("").astype(str).str.upper().apply(
                lambda value: any(keyword in value for keyword in keywords)
            )
            rows = remaining[mask]
            remaining = remaining[~mask]
        else:
            rows = remaining

        if rows.empty:
            continue

        sheet.merge_cells(start_row=row_number, start_column=1, end_row=row_number, end_column=len(export_columns))
        cell = sheet.cell(row=row_number, column=1, value=title)
        cell.fill, cell.font, cell.alignment = fill_section, white_bold, center
        row_number += 1

        for column_index, (_, heading) in enumerate(export_columns, 1):
            cell = sheet.cell(row=row_number, column=column_index, value=heading)
            cell.fill, cell.font, cell.alignment, cell.border = fill_header, white_bold, center, border
        row_number += 1

        for _, data in rows.iterrows():
            for column_index, (key, _) in enumerate(export_columns, 1):
                value = data.get(key, "")
                if pd.isna(value):
                    value = ""
                cell = sheet.cell(row=row_number, column=column_index, value=value)
                cell.fill, cell.font, cell.alignment, cell.border = fill_data, black_font, left, border
            row_number += 1
        row_number += 1

    widths = [7, 11, 24, 7, 10, 28, 17, 17, 9, 17, 18, 28, 18, 18, 10, 22]
    for index, width in enumerate(widths, 1):
        sheet.column_dimensions[chr(64 + index)].width = width
    sheet.freeze_panes = "A3"

    output = BytesIO()
    workbook.save(output)
    output.seek(0)
    return output


def exportar_reporte_excel(data: dict[str, pd.DataFrame], report_date: datetime) -> BytesIO:
    from openpyxl import Workbook
    from openpyxl.styles import Alignment, Font, PatternFill

    workbook = Workbook()
    overview = workbook.active
    overview.title = "Resumen"
    cyan = PatternFill("solid", fgColor="00B0F0")
    dark = PatternFill("solid", fgColor="123047")
    white = Font(color="FFFFFF", bold=True)

    overview.merge_cells("A1:D1")
    overview["A1"] = f"REPORTE DE OCUPACIÃ“N â€” {report_date.strftime('%B %d, %Y').upper()}"
    overview["A1"].fill, overview["A1"].font = cyan, Font(color="FFFFFF", bold=True, size=14)
    overview["A1"].alignment = Alignment(horizontal="center")
    overview.append([])
    overview.append(["MÃ©trica", "Reservas", "VIPs", "Habitaciones"])
    for cell in overview[3]:
        cell.fill, cell.font = dark, white

    for title, frame in data.items():
        vip_count = frame["info"].fillna("").astype(str).str.upper().str.contains("VIP").sum()
        rooms = ", ".join(frame["room"].dropna().astype(str).replace("", pd.NA).dropna().tolist()) or "â€”"
        overview.append([title, len(frame), int(vip_count), rooms])

        sheet_name = title[:31]
        sheet = workbook.create_sheet(sheet_name)
        visible = frame[[column for column in DISPLAY_COLUMNS if column in frame.columns]].copy()
        visible.to_excel(sheet, index=False)
        for cell in sheet[1]:
            cell.fill, cell.font = dark, white
        sheet.freeze_panes = "A2"

    for column, width in zip(("A", "B", "C", "D"), (30, 12, 12, 55)):
        overview.column_dimensions[column].width = width

    output = BytesIO()
    workbook.save(output)
    output.seek(0)
    return output


# -----------------------------------------------------------------------------
# Interfaz comÃºn
# -----------------------------------------------------------------------------

def show_header() -> None:
    import streamlit.components.v1 as components

    header_left, header_right = st.columns([1.6, 1])
    with header_left:
        st.markdown(
            """
<div class="page-title">
  <div class="brand">
    <div class="brand-mark">WA</div>
    <div>
      <div class="brand-name">WALDORF ASTORIA</div>
      <div class="brand-place">COSTA RICA Â· PUNTA CACIQUE</div>
      <div class="brand-product">Concierge Master <span style="color:#6f879a">v5.1</span></div>
    </div>
  </div>
</div>
            """,
            unsafe_allow_html=True,
        )
    with header_right:
        components.html(
            """
<!doctype html>
<html>
<head>
<style>
  body { margin:0; background:transparent; font-family:Segoe UI,sans-serif; text-align:right; }
  #local-clock { color:#00e5ff; font-size:16px; font-weight:800; line-height:1.2; text-shadow:0 0 12px rgba(0,229,255,.45); }
  #local-date { color:#00e5ff; font-size:14px; font-weight:700; margin-top:5px; text-transform:capitalize; }
</style>
</head>
<body>
  <div id="local-clock">--:--:--</div>
  <div id="local-date">Loading dateâ€¦</div>
<script>
  const clock = document.getElementById('local-clock');
  const date = document.getElementById('local-date');
  const timeFormatter = new Intl.DateTimeFormat(undefined, { hour: 'numeric', minute: '2-digit', second: '2-digit' });
  const dateFormatter = new Intl.DateTimeFormat(undefined, { weekday: 'long', month: 'long', day: 'numeric', year: 'numeric' });
  function updateLocalClock() {
    const now = new Date();
    clock.textContent = timeFormatter.format(now);
    date.textContent = dateFormatter.format(now);
  }
  updateLocalClock();
  setInterval(updateLocalClock, 1000);
</script>
</body>
</html>
            """,
            height=54,
            scrolling=False,
        )


def render_quick_links() -> None:
    links = "".join(
        f'<a class="quick-link" href="{html.escape(url, quote=True)}" style="background:{color}">{label}</a>'
        for label, url, color in QUICK_LINKS
    )
    st.markdown(f'<div class="quick-links">{links}</div>', unsafe_allow_html=True)


def render_action_links() -> None:
    actions = [
        ("NUEVA", "nueva", "#00B8D4", "#041319"),
        ("IMPORTAR", "importar", "#16A34A", "#FFFFFF"),
        ("EXPORTAR", "exportar", "#2563EB", "#FFFFFF"),
        ("REPORTE", "reporte", "#D97706", "#1C1300"),
        ("AGENDA", "agenda", "#7C3AED", "#FFFFFF"),
        ("CALCULADORA", "calculadora", "#E11D48", "#FFFFFF"),
    ]
    buttons = "".join(
        f'<a class="action-link" href="{url_with(action=action)}" style="background:{color};color:{text_color} !important">{label}</a>'
        for label, action, color, text_color in actions
    )
    st.markdown(f'<div class="action-links">{buttons}</div>', unsafe_allow_html=True)


def apply_filters(df: pd.DataFrame) -> tuple[pd.DataFrame, dict[str, str]]:
    checkout = str(st.query_params.get("checkout_filtro", ""))
    arrival_day = str(st.query_params.get("fecha_date", ""))
    search = st.session_state.get("global_search", "").strip()
    filters: dict[str, str] = {}
    result = df.copy()

    if checkout:
        result = result[result["check_out"] == checkout]
        filters["checkout"] = checkout
    if arrival_day:
        selected = date_from_filter(arrival_day)
        if selected:
            formatted = selected.strftime("%B %d, %Y")
            result = result[result["check_in"] == formatted]
            filters["arrival"] = formatted
    if search:
        text = search.lower()
        mask = result.astype(str).apply(
            lambda row: row.str.lower().str.contains(text, na=False).any(), axis=1
        )
        result = result[mask]
        filters["search"] = search
    return result, filters


def render_category_chart(df: pd.DataFrame) -> None:
    counts: list[tuple[str, int, str]] = []
    info = df["info"].fillna("").astype(str).str.upper() if "info" in df.columns else pd.Series(dtype=str)
    for category, color in CATEGORY_COLORS.items():
        counts.append((category, int(info.str.contains(category, na=False).sum()), color))
    categorized = sum(value for _, value, _ in counts)
    counts.append(("LEISURE", max(0, len(df) - categorized), "#3B82F6"))
    maximum = max((value for _, value, _ in counts), default=1) or 1

    content = '<div class="panel"><div class="panel-title">Guest categories</div>'
    for label, value, color in sorted(counts, key=lambda item: item[1], reverse=True):
        width = value / maximum * 100
        content += (
            f'<div class="category-row"><div class="category-label">{label}</div>'
            f'<div class="category-track"><div class="category-fill" style="width:{width:.1f}%;background:{color}"></div></div>'
            f'<div class="category-value">{value}</div></div>'
        )
    st.markdown(content + "</div>", unsafe_allow_html=True)


# -----------------------------------------------------------------------------
# Vistas secundarias
# -----------------------------------------------------------------------------

def render_back_link() -> None:
    st.markdown(
        '<a class="action-link" href="?" style="background:#334155;max-width:185px">â† REGRESAR A LA TABLA</a>',
        unsafe_allow_html=True,
    )


def render_new_reservation() -> None:
    st.subheader("Nueva reservaciÃ³n")
    render_back_link()
    with st.form("new_reservation", clear_on_submit=True):
        first = st.columns(4)
        eta = first[0].text_input("ETA", value=datetime.now().strftime("%I:%M %p").lstrip("0"))
        name = first[1].text_input("Nombre *")
        qty = first[2].number_input("HuÃ©spedes", min_value=0, value=1)
        room = first[3].text_input("HabitaciÃ³n")
        second = st.columns(4)
        email = second[0].text_input("Email")
        check_in = second[1].date_input("Check-in")
        check_out = second[2].date_input("Check-out", value=datetime.now().date() + timedelta(days=1))
        res_number = second[3].text_input("Reservation #")
        third = st.columns(4)
        phone = third[0].text_input("TelÃ©fono")
        info = third[1].text_input("Information")
        ird = third[2].text_input("IRD")
        hsk = third[3].text_input("HSK")
        fourth = st.columns(2)
        rate = fourth[0].text_input("Rate")
        trans = fourth[1].text_input("Transportation")
        submitted = st.form_submit_button("GUARDAR RESERVACIÃ“N", use_container_width=True)

    if submitted:
        if not name.strip():
            st.error("El nombre del huÃ©sped es obligatorio.")
            return
        if check_out < check_in:
            st.error("La fecha de check-out no puede ser anterior al check-in.")
            return
        insertar_reserva({
            "eta": eta.strip(), "name": name.strip(), "qty": int(qty), "room": room.strip(),
            "email": email.strip(), "check_in": check_in.strftime("%B %d, %Y"),
            "check_out": check_out.strftime("%B %d, %Y"), "res_number": res_number.strip(),
            "phone": phone.strip(), "info": info.strip(), "ird": ird.strip(), "hsk": hsk.strip(),
            "rate": rate.strip(), "trans": trans.strip(),
        })
        st.success("ReservaciÃ³n guardada correctamente.")
        st.query_params.clear()
        st.rerun()


def render_edit_reservation() -> None:
    reservation = st.session_state.get("selected_reservation")
    if not reservation:
        st.error("Selecciona una reserva de la tabla antes de editar.")
        render_back_link()
        return

    st.subheader(f"Editar reservaciÃ³n Â· {safe_text(reservation.get('name', ''))}")
    render_back_link()
    check_in_default = parse_fecha(reservation.get("check_in")) or datetime.now()
    check_out_default = parse_fecha(reservation.get("check_out")) or datetime.now() + timedelta(days=1)
    qty_default = pd.to_numeric(reservation.get("qty", 0), errors="coerce")
    qty_default = 0 if pd.isna(qty_default) else int(qty_default)

    with st.form("edit_reservation"):
        first = st.columns(4)
        eta = first[0].text_input("ETA", value=str(reservation.get("eta", "")))
        name = first[1].text_input("Nombre *", value=str(reservation.get("name", "")))
        qty = first[2].number_input("HuÃ©spedes", min_value=0, value=qty_default)
        room = first[3].text_input("HabitaciÃ³n", value=str(reservation.get("room", "")))
        second = st.columns(4)
        email = second[0].text_input("Email", value=str(reservation.get("email", "")))
        check_in = second[1].date_input("Check-in", value=check_in_default.date())
        check_out = second[2].date_input("Check-out", value=check_out_default.date())
        res_number = second[3].text_input("Reservation #", value=str(reservation.get("res_number", "")))
        third = st.columns(4)
        phone = third[0].text_input("TelÃ©fono", value=str(reservation.get("phone", "")))
        info = third[1].text_input("Information", value=str(reservation.get("info", "")))
        ird = third[2].text_input("IRD", value=str(reservation.get("ird", "")))
        hsk = third[3].text_input("HSK", value=str(reservation.get("hsk", "")))
        fourth = st.columns(2)
        rate = fourth[0].text_input("Rate", value=str(reservation.get("rate", "")))
        trans = fourth[1].text_input("Transportation", value=str(reservation.get("trans", "")))
        submitted = st.form_submit_button("GUARDAR CAMBIOS", use_container_width=True)

    if submitted:
        if not name.strip():
            st.error("El nombre del huÃ©sped es obligatorio.")
            return
        if check_out < check_in:
            st.error("La fecha de check-out no puede ser anterior al check-in.")
            return
        actualizar_reserva(reservation["id"], {
            "eta": eta.strip(), "name": name.strip(), "qty": int(qty), "room": room.strip(),
            "email": email.strip(), "check_in": check_in.strftime("%B %d, %Y"),
            "check_out": check_out.strftime("%B %d, %Y"), "res_number": res_number.strip(),
            "phone": phone.strip(), "info": info.strip(), "ird": ird.strip(), "hsk": hsk.strip(),
            "rate": rate.strip(), "trans": trans.strip(),
        })
        st.success("Reserva actualizada correctamente.")
        clear_page()


def render_import() -> None:
    st.subheader("Importar reservaciones desde Excel")
    render_back_link()
    st.info("Columnas requeridas: " + ", ".join(IMPORT_COLUMNS))
    uploaded = st.file_uploader("Archivo Excel", type=["xlsx", "xls"])
    if not uploaded:
        return

    try:
        frame = pd.read_excel(uploaded)
    except Exception as exc:
        st.error(f"No se pudo leer el archivo: {exc}")
        return

    frame.columns = [str(column).strip().lower() for column in frame.columns]
    missing = [column for column in IMPORT_COLUMNS if column not in frame.columns]
    if missing:
        st.error("Faltan estas columnas: " + ", ".join(missing))
        return

    preview = frame[IMPORT_COLUMNS].copy()
    preview["check_in"] = preview["check_in"].map(normalizar_fecha)
    preview["check_out"] = preview["check_out"].map(normalizar_fecha)
    st.success(f"Archivo vÃ¡lido: {len(preview)} reservaciones detectadas.")
    st.dataframe(preview, use_container_width=True, hide_index=True, height=300)

    if st.button("IMPORTAR A BASE DE DATOS", type="primary", use_container_width=True):
        records: list[dict] = []
        for _, row in preview.iterrows():
            record = {}
            for column in IMPORT_COLUMNS:
                value = row[column]
                if pd.isna(value):
                    value = ""
                elif column == "qty":
                    value = int(pd.to_numeric(value, errors="coerce") or 0)
                else:
                    value = str(value).strip()
                record[column] = value
            records.append(record)
        try:
            insertar_lote(records)
            st.success(f"{len(records)} reservaciones importadas correctamente.")
        except Exception as exc:
            st.error(f"No se completÃ³ la importaciÃ³n: {exc}")


def render_export(df: pd.DataFrame) -> None:
    st.subheader("Exportar reservaciones a Excel")
    render_back_link()
    filtered, _ = apply_filters(df)
    st.info(f"Se exportarÃ¡n {len(filtered)} reservaciones, organizadas por categorÃ­a.")
    st.dataframe(filtered[DISPLAY_COLUMNS], use_container_width=True, hide_index=True, height=300)
    if not filtered.empty:
        st.download_button(
            "DESCARGAR EXCEL",
            data=exportar_excel_por_categorias(filtered),
            file_name=f"Arrivals_{datetime.now():%Y%m%d_%H%M}.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            use_container_width=True,
        )


def render_agenda(df: pd.DataFrame) -> None:
    st.subheader("Agenda de reservaciones")
    render_back_link()
    selected_date = st.date_input("Fecha", value=datetime.now().date())
    formatted = selected_date.strftime("%B %d, %Y")
    arrivals = df[df["check_in"] == formatted]
    departures = df[df["check_out"] == formatted]
    c1, c2, c3 = st.columns(3)
    c1.metric("Llegan", len(arrivals))
    c2.metric("Salen", len(departures))
    c3.metric("Total en base", len(df))
    arrivals_tab, departures_tab = st.tabs(["Llegadas", "Salidas"])
    with arrivals_tab:
        st.dataframe(arrivals[[column for column in DISPLAY_COLUMNS if column in arrivals]], use_container_width=True, hide_index=True)
    with departures_tab:
        st.dataframe(departures[[column for column in DISPLAY_COLUMNS if column in departures]], use_container_width=True, hide_index=True)


def render_report(df: pd.DataFrame) -> None:
    st.subheader("Reporte de ocupaciÃ³n diario")
    render_back_link()
    today = datetime.now()
    tomorrow = today + timedelta(days=1)
    today_label, tomorrow_label = today.strftime("%B %d, %Y"), tomorrow.strftime("%B %d, %Y")

    parsed_check_in = df["check_in"].map(parse_fecha)
    parsed_check_out = df["check_out"].map(parse_fecha)
    in_house = df[(parsed_check_in <= today) & (parsed_check_out > today)]
    departures_today = df[df["check_out"] == today_label]
    arrivals_today = df[df["check_in"] == today_label]
    departures_tomorrow = df[df["check_out"] == tomorrow_label]
    arrivals_tomorrow = df[df["check_in"] == tomorrow_label]

    report_data = {
        "En casa": in_house,
        "Salen hoy": departures_today,
        "Salen maÃ±ana": departures_tomorrow,
        "Llegan hoy": arrivals_today,
        "Llegan maÃ±ana": arrivals_tomorrow,
    }
    metrics = st.columns(5)
    for column, (label, frame) in zip(metrics, report_data.items()):
        column.metric(label, len(frame))
    st.markdown("<div style='height:8px'></div>", unsafe_allow_html=True)
    st.dataframe(
        pd.DataFrame([
            {"CategorÃ­a": label, "Reservas": len(frame), "Habitaciones": ", ".join(frame["room"].dropna().astype(str).tolist()) or "â€”"}
            for label, frame in report_data.items()
        ]),
        use_container_width=True,
        hide_index=True,
    )
    st.download_button(
        "DESCARGAR REPORTE EXCEL",
        data=exportar_reporte_excel(report_data, today),
        file_name=f"Reporte_Ocupacion_{today:%Y%m%d_%H%M}.xlsx",
        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        use_container_width=True,
    )


def render_calculator() -> None:
    st.subheader("Calculadora")
    render_back_link()
    st.components.v1.html(
        """
<!doctype html><html><head><style>
body{margin:0;background:#080b12;font-family:Segoe UI,sans-serif;display:grid;place-items:center;padding:12px;color:#eafaff}
.calculator{width:300px;padding:16px;background:#101827;border:1px solid #284057;border-radius:16px;box-shadow:0 12px 35px #0008}
#display{background:#061c2b;border:1px solid #00e5ff;border-radius:10px;color:#00e5ff;font:700 30px monospace;padding:14px;text-align:right;overflow:hidden;margin-bottom:11px}
.grid{display:grid;grid-template-columns:repeat(4,1fr);gap:7px}button{border:0;border-radius:8px;padding:13px 6px;background:#203449;color:#eafaff;font-weight:800;font-size:16px;cursor:pointer}button:hover{filter:brightness(1.2)}.op{background:#7c3aed}.equal{background:#00c6df;color:#01171d}.clear{background:#e11d48}
</style></head><body><div class="calculator"><div id="display">0</div><div class="grid">
<button class="clear" onclick="clearAll()">C</button><button onclick="backspace()">âŒ«</button><button class="op" onclick="add('/')">Ã·</button><button class="op" onclick="add('*')">Ã—</button>
<button onclick="add('7')">7</button><button onclick="add('8')">8</button><button onclick="add('9')">9</button><button class="op" onclick="add('-')">âˆ’</button>
<button onclick="add('4')">4</button><button onclick="add('5')">5</button><button onclick="add('6')">6</button><button class="op" onclick="add('+')">+</button>
<button onclick="add('1')">1</button><button onclick="add('2')">2</button><button onclick="add('3')">3</button><button onclick="add('.')">.</button>
<button onclick="add('0')">0</button><button onclick="add('%')">%</button><button onclick="toggleSign()">Â±</button><button class="equal" onclick="calculate()">=</button>
</div></div><script>
let value='';const out=document.getElementById('display');function draw(){out.textContent=value||'0'}function add(v){if('0123456789.'.includes(v)&&out.textContent==='Error')value='';value+=v;draw()}function clearAll(){value='';draw()}function backspace(){value=value.slice(0,-1);draw()}function toggleSign(){value=value.startsWith('-')?value.slice(1):'-'+value;draw()}function calculate(){try{if(!/^[0-9+*/.()% -]+$/.test(value))throw Error();value=String(Function('return ('+value+')')());draw()}catch(e){value='';out.textContent='Error'}}document.addEventListener('keydown',e=>{if('0123456789.+-*/%'.includes(e.key))add(e.key);else if(e.key==='Enter')calculate();else if(e.key==='Backspace')backspace();else if(e.key==='Escape')clearAll()});
</script></body></html>
        """,
        height=420,
    )


def render_letter() -> None:
    reservation = st.session_state.get("selected_reservation")
    st.subheader("Carta de despedida")
    render_back_link()
    if not reservation:
        st.error("Selecciona una reserva de la tabla antes de crear la carta.")
        return

    guest_name = str(reservation.get("name", "")).strip()
    if not guest_name:
        st.error("La reserva seleccionada no tiene un nombre de huÃ©sped.")
        return

    template_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "plantilla_despedida.docx")
    if not os.path.exists(template_path):
        st.warning("No se encontrÃ³ `plantilla_despedida.docx` junto a este archivo.")
        st.info("Crea esa plantilla Word y usa `{{NOMBRE}}` donde debe aparecer el nombre del huÃ©sped.")
        return

    try:
        from docx import Document
    except ImportError:
        st.error("Falta `python-docx`. InstÃ¡lalo con el archivo requirements_streamlit.txt actualizado.")
        return

    try:
        document = Document(template_path)
        replacements = 0

        def replace_paragraph(paragraph) -> None:
            nonlocal replacements
            if "{{NOMBRE}}" not in paragraph.text:
                return
            text = paragraph.text.replace("{{NOMBRE}}", guest_name)
            for run in paragraph.runs:
                run.text = ""
            if paragraph.runs:
                paragraph.runs[0].text = text
            else:
                paragraph.add_run(text)
            replacements += 1

        for paragraph in document.paragraphs:
            replace_paragraph(paragraph)
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        replace_paragraph(paragraph)

        output = BytesIO()
        document.save(output)
        output.seek(0)
        safe_name = "".join(char if char.isalnum() or char in "_-" else "_" for char in guest_name)
        st.success(f"Carta preparada para {guest_name}. Placeholders sustituidos: {replacements}.")
        st.download_button(
            "DESCARGAR CARTA DE DESPEDIDA",
            data=output,
            file_name=f"Carta_Despedida_{safe_name}_{datetime.now():%Y%m%d_%H%M}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
        )
    except Exception as exc:
        st.error(f"No se pudo generar la carta: {exc}")


def render_delete() -> None:
    reservation = st.session_state.get("selected_reservation")
    st.subheader("Eliminar reservaciÃ³n")
    render_back_link()
    if not reservation:
        st.error("Selecciona una reserva antes de solicitar el borrado.")
        return
    st.warning(f"Se eliminarÃ¡ permanentemente la reserva de {reservation.get('name', 'este huÃ©sped')}.")
    with st.form("delete_reservation"):
        password = st.text_input("Clave de autorizaciÃ³n", type="password")
        confirmed = st.form_submit_button("CONFIRMAR Y BORRAR", type="primary")
    if confirmed:
        expected_password = st.secrets.get("DELETE_PASSWORD", "")
        if not expected_password:
            st.error("Configura `DELETE_PASSWORD` en los Secrets antes de habilitar el borrado.")
        elif password != expected_password:
            st.error("Clave incorrecta.")
        else:
            eliminar_reserva(reservation["id"])
            st.success("Reserva eliminada correctamente.")
            clear_page()


# -----------------------------------------------------------------------------
# Tabla AG Grid
# -----------------------------------------------------------------------------

CATEGORY_CELL_STYLE = JsCode(
    """
function(params) {
  const source = String((params.data && params.data.info) || '').toUpperCase();
  const current = String(params.value || '').toUpperCase();
  const haystack = source + ' ' + current;
  const categories = [
    ['VIP', '#00E5FF'], ['BIRTHDAY', '#FF5252'], ['HONEYMOON', '#FF9800'],
    ['BABYMOON', '#A78BFA'], ['ANNIVERSARY', '#4ADE80'], ['RELAXURY', '#F472B6'],
    ['TEAM MEMBER', '#FACC15']
  ];
  for (const [keyword, color] of categories) {
    if (haystack.includes(keyword)) {
      return { backgroundColor: color + '24', color: color, fontWeight: '800' };
    }
  }
  return null;
}
"""
)

VIP_ROW_STYLE = JsCode(
    """
function(params) {
  if (String((params.data && params.data.info) || '').toUpperCase().includes('VIP')) {
    return { backgroundColor: '#191507', color: '#f6d670' };
  }
  return null;
}
"""
)

GRID_CSS = {
    ".ag-root-wrapper": {
        "border": "1px solid #1f384d !important",
        "border-radius": "10px !important",
        "overflow": "hidden !important",
        "background-color": "#0b111b !important",
    },
    ".ag-header": {
        "background-color": "#00E5FF !important",
        "border-bottom": "none !important",
    },
    ".ag-header-cell": {"border-right": "none !important"},
    ".ag-header-cell-label": {"color": "#00151d !important", "font-weight": "900 !important", "letter-spacing": ".25px"},
    ".ag-header-cell-text": {"color": "#00151d !important"},
    ".ag-row": {"background-color": "#0b111b !important", "border-bottom": "none !important"},
    ".ag-row-odd": {"background-color": "#0e1723 !important"},
    ".ag-cell": {"border-right": "none !important", "border-bottom": "none !important", "color": "#e6f3fb !important", "font-size": "12px"},
    ".ag-row-hover": {"background-color": "#132a3a !important"},
    ".ag-row-selected": {"background-color": "#00E5FF !important"},
    ".ag-row-selected .ag-cell": {"color": "#00151d !important", "font-weight": "800 !important"},
    ".ag-paging-panel": {"border-top": "none !important", "background-color": "#0b111b !important", "color": "#cceaf6 !important"},
}


def render_reservations_grid(df: pd.DataFrame) -> None:
    visible = df[[column for column in DISPLAY_COLUMNS if column in df.columns]].copy()
    builder = GridOptionsBuilder.from_dataframe(visible)
    builder.configure_default_column(resizable=True, sortable=True, filter=True, minWidth=80)
    builder.configure_selection(selection_mode="single", use_checkbox=False)
    builder.configure_grid_options(
        getRowStyle=VIP_ROW_STYLE,
        rowHeight=34,
        headerHeight=37,
        suppressCellFocus=True,
        animateRows=True,
    )

    fields = {
        "eta": ("ETA", 85), "name": ("NAME", 165), "qty": ("QTY", 65), "room": ("ROOM", 75),
        "email": ("EMAIL", 200), "check_in": ("CHECK IN", 130), "check_out": ("CHECK OUT", 130),
        "nights": ("NOCHES", 85), "res_number": ("RESERVATION", 130), "phone": ("PHONE", 135),
        "info": ("INFORMATION", 210), "ird": ("IRD", 150), "hsk": ("HSK", 120),
        "rate": ("RATE", 85), "trans": ("TRANS", 170),
    }
    for field, (header, width) in fields.items():
        if field not in visible.columns:
            continue
        config: dict = {"header_name": header, "width": width}
        if field in {"info", "ird", "trans"}:
            config["cellStyle"] = CATEGORY_CELL_STYLE
        if field == "nights":
            config["type"] = ["numericColumn"]
        builder.configure_column(field, **config)

    response = AgGrid(
        visible,
        gridOptions=builder.build(),
        custom_css=GRID_CSS,
        theme="alpine",
        height=625,
        fit_columns_on_grid_load=False,
        allow_unsafe_jscode=True,
        update_mode=GridUpdateMode.SELECTION_CHANGED,
        key="concierge_reservations_grid",
    )

    selected = response.get("selected_rows", [])
    if isinstance(selected, pd.DataFrame):
        selected = selected.to_dict("records")
    if selected:
        selected_id = selected[0].get("id")
        # AgGrid recibe una vista sin `id`; la buscamos con una combinaciÃ³n estable de datos.
        candidate = df[
            (df["name"].astype(str) == str(selected[0].get("name", "")))
            & (df["res_number"].astype(str) == str(selected[0].get("res_number", "")))
            & (df["check_in"].astype(str) == str(selected[0].get("check_in", "")))
        ]
        if not candidate.empty:
            row = candidate.iloc[0].to_dict()
            if st.session_state.get("selected_reservation_id") != row.get("id"):
                st.session_state["selected_reservation"] = row
                st.session_state["selected_reservation_id"] = row.get("id")
                st.rerun()


# -----------------------------------------------------------------------------
# Dashboard principal
# -----------------------------------------------------------------------------

def render_dashboard(df: pd.DataFrame) -> None:
    st.markdown(f'<div class="total-strip">TOTAL RESERVAS <strong>{len(df)}</strong></div>', unsafe_allow_html=True)
    st.markdown("<div style='height:8px'></div>", unsafe_allow_html=True)

    left, right = st.columns([1.35, 5.65], gap="medium")
    with left:
        st.markdown('<div class="panel"><div class="panel-title">Checking out rooms</div>', unsafe_allow_html=True)
        today = datetime.now()
        checkout_links = []
        for offset in range(8):
            date = today + timedelta(days=offset)
            stored = date.strftime("%B %d, %Y")
            count = int((df["check_out"] == stored).sum())
            checkout_links.append(
                f'<a class="quick-link" href="{url_with(checkout_filtro=stored)}" style="background:#{"0F766E" if offset % 2 == 0 else "155E75"}">{date:%d-%b}: {count}</a>'
            )
        st.markdown('<div style="display:grid;grid-template-columns:repeat(2,1fr);gap:6px">' + "".join(checkout_links) + "</div>", unsafe_allow_html=True)
        st.markdown('<div style="margin-top:8px"><a class="action-link" href="?" style="background:#334155">VER TODAS</a></div>', unsafe_allow_html=True)
        st.markdown("</div>", unsafe_allow_html=True)

        filter_default = date_from_filter(str(st.query_params.get("fecha_date", "")))
        selected_date = st.date_input("Filtrar por check-in", value=(filter_default or datetime.now()).date(), key="arrival_date")
        date_link = url_with(fecha_date=selected_date.strftime("%Y-%m-%d"))
        st.markdown(
            f'<a class="action-link" href="{date_link}" style="background:#0891B2;margin-top:6px">APLICAR FECHA</a>',
            unsafe_allow_html=True,
        )
        st.markdown('<a class="action-link" href="?" style="background:#475569;margin-top:6px">LIMPIAR FILTROS</a>', unsafe_allow_html=True)

    with right:
        search_left, search_right = st.columns([1, 4])
        search_left.markdown("<div style='padding-top:8px;color:#8ca4ba;font-size:12px;text-align:right'>BÃºsqueda rÃ¡pida</div>", unsafe_allow_html=True)
        search_right.text_input(
            "BÃºsqueda rÃ¡pida",
            key="global_search",
            placeholder="Nombre, telÃ©fono, reserva, categorÃ­aâ€¦",
            label_visibility="collapsed",
        )
        filtered, filters = apply_filters(df)
        render_category_chart(filtered)
        relaxury = int(filtered.astype(str).apply(lambda column: column.str.upper().str.contains("RELAXURY", na=False)).any(axis=1).sum())
        st.markdown(f'<div class="total-strip" style="border-color:#f472b6">RELAXURY <strong style="color:#f472b6">{relaxury}</strong></div>', unsafe_allow_html=True)
        render_action_links()
        render_quick_links()

    if filters:
        captions = []
        if "checkout" in filters:
            captions.append("Check-out: " + safe_text(filters["checkout"]))
        if "arrival" in filters:
            captions.append("Check-in: " + safe_text(filters["arrival"]))
        if "search" in filters:
            captions.append("BÃºsqueda: " + safe_text(filters["search"]))
        st.caption(" Â· ".join(captions))

    st.markdown("<div style='height:8px'></div>", unsafe_allow_html=True)
    render_reservations_grid(filtered)

    selected = st.session_state.get("selected_reservation")
    if selected:
        name, room = safe_text(selected.get("name", "N/A")), safe_text(selected.get("room", "â€”"))
        st.markdown(f'<div class="selection-banner"><b>â— RESERVA SELECCIONADA</b> &nbsp; {name} &nbsp;|&nbsp; Room: {room}</div>', unsafe_allow_html=True)
        st.markdown(
            '<div class="action-links" style="grid-template-columns:repeat(4,minmax(110px,1fr));max-width:700px">'
            f'<a class="action-link" href="{url_with(action="editar")}" style="background:#D97706">EDITAR</a>'
            f'<a class="action-link" href="{url_with(action="carta")}" style="background:#7C3AED">CARTA</a>'
            f'<a class="action-link" href="{url_with(action="cancelar")}" style="background:#E11D48">BORRAR</a>'
            '<a class="action-link" href="?" style="background:#475569">DESELECCIONAR</a>'
            '</div>',
            unsafe_allow_html=True,
        )


# -----------------------------------------------------------------------------
# App
# -----------------------------------------------------------------------------

show_header()
reservations = cargar_reservaciones()
action = get_action()

if action == "nueva":
    render_new_reservation()
elif action == "editar":
    render_edit_reservation()
elif action == "importar":
    render_import()
elif action == "exportar":
    render_export(reservations)
elif action == "agenda":
    render_agenda(reservations)
elif action == "reporte":
    render_report(reservations)
elif action == "calculadora":
    render_calculator()
elif action == "carta":
    render_letter()
elif action == "cancelar":
    render_delete()
else:
    render_dashboard(reservations)
